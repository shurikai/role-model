from datetime import datetime

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Twips
from docx.text.paragraph import Paragraph

from models import EducationEntry, EmployerBlock, Identity, ProjectEntry, Resume

FONT_NAME = "Arial"
ACCENT_COLOR = RGBColor(0x28, 0x01, 0x37)
TENURE_COLOR = RGBColor(0x59, 0x59, 0x59)

# Page geometry (dxa/twips, 1/20 pt). CONTENT_WIDTH_DXA is the tab stop
# position for role/project header lines -- hardcoded per the canonical
# formatting spec rather than derived from margins at render time.
MARGIN_DXA = 1080
CONTENT_WIDTH_DXA = 10240

SECTION_HEADING_SPACE_BEFORE = Twips(320)
SECTION_HEADING_SPACE_AFTER = Twips(20)
RULE_SPACE_BEFORE = Twips(0)
RULE_SPACE_AFTER = Twips(120)

HEADER_SPACE_BEFORE = Twips(160)

TITLE_SPACE_BEFORE = Twips(0)
TITLE_SPACE_AFTER = Twips(60)

BULLET_INDENT_LEFT = Twips(480)
BULLET_INDENT_HANGING = Twips(280)
BULLET_SPACE_BEFORE = Twips(20)
BULLET_SPACE_AFTER = Twips(20)


def build_resume_document(resume: Resume) -> DocumentObject:
    doc = Document()
    _configure_page(doc)
    _configure_base_style(doc)

    _render_identity(doc, resume.identity)
    _render_summary(doc, resume.summary)
    _render_skills(doc, resume.skills)
    _render_experience(doc, resume.experience)
    _render_projects(doc, resume.projects)
    _render_education(doc, resume.education)

    return doc


def _configure_page(doc: DocumentObject) -> None:
    section = doc.sections[0]
    section.top_margin = Twips(MARGIN_DXA)
    section.bottom_margin = Twips(MARGIN_DXA)
    section.left_margin = Twips(MARGIN_DXA)
    section.right_margin = Twips(MARGIN_DXA)


def _configure_base_style(doc: DocumentObject) -> None:
    """Sets Arial as the document font and zeroes the Normal style's default
    paragraph spacing. Word's built-in Normal style carries ~8pt space-after
    by default; left alone, every paragraph we don't explicitly format would
    inherit it and silently reintroduce the bloat this refactor removes."""
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT_NAME)

    style.paragraph_format.space_before = Twips(0)
    style.paragraph_format.space_after = Twips(0)


def _set_bottom_border(paragraph: Paragraph, color: str, size: int = 6) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _section_heading(doc: DocumentObject, text: str) -> None:
    """Section heading (SUMMARY, SKILLS, EXPERIENCE, ...) with an accent-color
    rule line below it. Explicit paragraph-level spacing, not a Word built-in
    Heading style -- built-in heading styles carry spacing well beyond what's
    needed and inflate page count independent of content.

    keep_with_next is set on both the heading and rule paragraphs so the
    heading, its rule line, and the first line of the section's content
    are never split across a page boundary. keep_with_next only binds a
    paragraph to the one immediately after it, so it has to be set on
    each link in the chain, not just the first."""
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = SECTION_HEADING_SPACE_BEFORE
    heading.paragraph_format.space_after = SECTION_HEADING_SPACE_AFTER
    heading.paragraph_format.keep_with_next = True
    run = heading.add_run(text)
    run.bold = True
    run.font.color.rgb = ACCENT_COLOR

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = RULE_SPACE_BEFORE
    rule.paragraph_format.space_after = RULE_SPACE_AFTER
    rule.paragraph_format.keep_with_next = True
    _set_bottom_border(rule, color="280137")


def _add_bullet(doc: DocumentObject, text: str) -> Paragraph:
    p = doc.add_paragraph(text, style="List Bullet")
    pf = p.paragraph_format
    pf.left_indent = BULLET_INDENT_LEFT
    pf.first_line_indent = -BULLET_INDENT_HANGING
    pf.space_before = BULLET_SPACE_BEFORE
    pf.space_after = BULLET_SPACE_AFTER
    return p


def _header_line(doc: DocumentObject, primary: str, secondary: str | None) -> None:
    """Two/three-run header line: primary (bold) + tab + secondary (gray),
    on one line via a hardcoded right tab stop. Per spec, the tab is its own
    run -- never embedded in the primary/secondary run text -- so the layout
    is driven by the tab stop, not by whitespace baked into a string.

    Carries its own space-before so consecutive role/project entries don't
    run together -- everything else in a header block (subtitle, bullets)
    is spaced tight against what precedes it, but the header line itself
    needs separation from the previous entry's last bullet.

    keep_with_next ties this line to the subtitle line immediately below
    it, so employer/date and title never land on opposite sides of a page
    break."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = HEADER_SPACE_BEFORE
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.tab_stops.add_tab_stop(
        Twips(CONTENT_WIDTH_DXA), WD_TAB_ALIGNMENT.RIGHT
    )

    primary_run = p.add_run(primary)
    primary_run.bold = True

    if secondary:
        p.add_run("\t")
        secondary_run = p.add_run(secondary)
        secondary_run.font.color.rgb = TENURE_COLOR


def _subtitle_line(doc: DocumentObject, text: str) -> None:
    """Italic subtitle line immediately below a header line (role title,
    project tagline) -- own paragraph, tight spacing, 2 lines total for the
    header block rather than 3.

    keep_with_next ties this line to whatever comes right after it (the
    first bullet, normally), so the header block's title never gets
    stranded from its own content. Bullets themselves are deliberately
    left free to break across a page boundary -- only the header chain
    needs protection."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = TITLE_SPACE_BEFORE
    p.paragraph_format.space_after = TITLE_SPACE_AFTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.italic = True


def _render_identity(doc: DocumentObject, identity: Identity) -> None:
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Twips(20)
    name_run = name_p.add_run(identity.name)
    name_run.bold = True
    name_run.font.size = Pt(16)

    contact_fields = [
        identity.headline,
        identity.location,
        identity.email,
        identity.phone,
        identity.linkedin_url,
        identity.github_url,
        identity.site_url,
    ]
    contact_line = " · ".join(field for field in contact_fields if field)
    if contact_line:
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_p.paragraph_format.space_after = SECTION_HEADING_SPACE_AFTER
        contact_p.add_run(contact_line)


def _render_summary(doc: DocumentObject, summary: str) -> None:
    _section_heading(doc, "SUMMARY")
    doc.add_paragraph(summary)


def _render_skills(doc: DocumentObject, skills: dict[str, list[str]]) -> None:
    _section_heading(doc, "SKILLS")
    for category, skills_list in skills.items():
        if skills_list:
            line = f"{category}: {', '.join(skills_list)}"
            _add_bullet(doc, line)


def _render_experience(doc: DocumentObject, experience: list[EmployerBlock]) -> None:
    _section_heading(doc, "EXPERIENCE")
    for employer_block in experience:
        for position in employer_block.positions:
            tenure = _format_tenure(position.started_on, position.ended_on)
            _header_line(doc, employer_block.employer, tenure)
            _subtitle_line(doc, position.title)

            for bullet in position.bullets:
                _add_bullet(doc, bullet.text)


def _render_projects(doc: DocumentObject, projects: list[ProjectEntry]) -> None:
    if not projects:
        return

    _section_heading(doc, "PROJECTS")
    for project in projects:
        tenure = None
        if project.started_on:
            ended = format_date(project.ended_on) if project.ended_on else "Present"
            tenure = f"{format_date(project.started_on)} – {ended}"
        _header_line(doc, project.name, tenure)

        if project.tagline:
            _subtitle_line(doc, project.tagline)

        for bullet in project.bullets:
            _add_bullet(doc, bullet.text)

        links = [project.writeup_url, project.repo_url, project.live_url]
        link_line = " · ".join(link for link in links if link)
        if link_line:
            doc.add_paragraph(link_line)


def _render_education(doc: DocumentObject, education: list[EducationEntry]) -> None:
    _section_heading(doc, "EDUCATION")
    for education_block in education:
        parts = [education_block.institution]
        if education_block.degree:
            parts.append(education_block.degree)
        if education_block.field_of_study:
            parts.append(education_block.field_of_study)
        if education_block.graduated:
            parts.append(education_block.graduated)
        doc.add_paragraph(" - ".join(parts))


def _format_tenure(started_on: str, ended_on: str | None) -> str:
    started = format_date(started_on)
    ended = format_date(ended_on) if ended_on else "Present"
    return f"{started} – {ended}"


def format_date(value: str) -> str:
    return datetime.strptime(value, "%Y-%m").strftime("%b %Y")
