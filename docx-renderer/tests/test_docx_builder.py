"""Document construction.

These tests assert two kinds of thing:

1. Nothing is silently dropped. The failure mode that matters here is a
   position or a bullet disappearing from the output while the render still
   returns 200 -- exactly the shape of the early-career dropout regression
   (#32) on the Go side, which was invisible until someone read the document.

2. The layout invariants CLAUDE.md calls load-bearing hold: no Word heading
   styles, keep_with_next on the header chains, and bullets deliberately left
   free to break across a page.

They do not assert on .docx bytes. Binary comparison against the tracked
fixtures would fail on every incidental spacing change and teach us nothing.
"""

from docx.document import Document as DocumentObject

from models import Resume
from renderer.docx_builder import build_resume_document


def paragraph_texts(doc: DocumentObject) -> list[str]:
    return [p.text for p in doc.paragraphs]


def document_text(doc: DocumentObject) -> str:
    return "\n".join(paragraph_texts(doc))


SECTION_HEADINGS = (
    "SUMMARY",
    "SKILLS",
    "EXPERIENCE",
    "PROJECTS",
    "EDUCATION",
    "CREDENTIALS",
)


def section_text(doc: DocumentObject, heading: str) -> str:
    """Text of one section only, heading exclusive.

    Whole-document substring checks are too loose for anything whose values
    also occur in prose: dropping the entire "Frameworks & Libraries" skill
    row still passed a document-wide search, because "Spring Boot" and "React"
    appear in bullet text elsewhere. Assertions about a section's contents have
    to be scoped to that section.
    """
    texts = paragraph_texts(doc)
    assert heading in texts, f"section not found: {heading}"
    start = texts.index(heading) + 1
    end = start
    while end < len(texts) and texts[end] not in SECTION_HEADINGS:
        end += 1
    return "\n".join(texts[start:end])


def test_builds_without_error(resume: Resume) -> None:
    doc = build_resume_document(resume)
    assert isinstance(doc, DocumentObject)
    assert len(doc.paragraphs) > 0


def test_identity_and_summary_are_rendered(resume: Resume) -> None:
    text = document_text(build_resume_document(resume))
    assert resume.identity.name in text
    assert resume.summary in text


def test_every_employer_appears(resume: Resume) -> None:
    text = document_text(build_resume_document(resume))
    for employer in resume.experience:
        assert employer.employer in text, f"employer dropped: {employer.employer}"


def test_every_position_title_appears(resume: Resume) -> None:
    text = document_text(build_resume_document(resume))
    for employer in resume.experience:
        for position in employer.positions:
            assert position.display_title in text, (
                f"position dropped: {position.display_title}"
            )


def test_normalized_role_is_rendered_in_place_of_the_verbatim_title(
    resume: Resume,
) -> None:
    """A verbatim title carries the employer's internal grade ladder rather
    than the job. "Programmer VII" reads as junior to anyone outside that
    company, and the database has carried the normalization all along."""
    position = resume.experience[0].positions[0]
    position.industry_role = "Senior Software Engineer"
    position.title = "Programmer VII"

    text = document_text(build_resume_document(resume))

    assert "Senior Software Engineer" in text
    assert "Programmer VII" not in text


def test_verbatim_title_is_the_fallback(resume: Resume) -> None:
    """Documents generated before industry_role existed must still render."""
    position = resume.experience[0].positions[0]
    position.industry_role = None
    position.title = "Programmer VII"

    assert "Programmer VII" in document_text(build_resume_document(resume))


def test_every_bullet_appears(resume: Resume) -> None:
    """The one that would have caught a dropped role."""
    text = document_text(build_resume_document(resume))
    for employer in resume.experience:
        for position in employer.positions:
            for bullet in position.bullets:
                assert bullet.text in text, f"bullet dropped: {bullet.text[:60]}"


def test_every_project_appears(resume: Resume) -> None:
    text = document_text(build_resume_document(resume))
    for project in resume.projects:
        assert project.name in text, f"project dropped: {project.name}"
        for bullet in project.bullets:
            assert bullet.text in text, f"project bullet dropped: {bullet.text[:60]}"


def test_every_skill_appears(resume: Resume) -> None:
    """Scoped to the SKILLS section on purpose -- see section_text."""
    skills_text = section_text(build_resume_document(resume), "SKILLS")
    for category, skills in resume.skills.items():
        assert category in skills_text, f"skill category dropped: {category}"
        for skill in skills:
            assert skill in skills_text, f"skill dropped: {category}/{skill}"


def test_education_appears(resume: Resume) -> None:
    if not resume.education:
        return
    education_text = section_text(build_resume_document(resume), "EDUCATION")
    for entry in resume.education:
        assert entry.institution in education_text


def test_expected_section_headings_present(resume: Resume) -> None:
    texts = paragraph_texts(build_resume_document(resume))
    assert "SUMMARY" in texts
    assert "SKILLS" in texts
    assert "EXPERIENCE" in texts


def test_projects_heading_tracks_whether_there_are_projects(resume: Resume) -> None:
    """An empty section must not leave a bare heading behind."""
    texts = paragraph_texts(build_resume_document(resume))
    assert ("PROJECTS" in texts) == bool(resume.projects)


def test_education_heading_tracks_whether_there_is_education(resume: Resume) -> None:
    texts = paragraph_texts(build_resume_document(resume))
    assert ("EDUCATION" in texts) == bool(resume.education)


def test_credentials_heading_tracks_whether_there_are_credentials(
    resume: Resume,
) -> None:
    texts = paragraph_texts(build_resume_document(resume))
    assert ("CREDENTIALS" in texts) == bool(resume.credentials)


def test_every_credential_appears(resume: Resume) -> None:
    """Credentials were modelled everywhere and rendered nowhere.

    The section existed in the schema, in Pydantic, and in the Go assembler,
    and build_resume_document simply never called a renderer for it -- so a
    certification that survived selection was dropped on the last step. For a
    licensed profession the licence is the most important line on the page.
    """
    if not resume.credentials:
        return
    text = section_text(build_resume_document(resume), "CREDENTIALS")
    for credential in resume.credentials:
        assert credential.name in text, f"credential dropped: {credential.name}"
        if credential.issuer:
            assert credential.issuer in text, f"issuer dropped: {credential.issuer}"


def test_all_optional_sections_empty_leaves_no_bare_headings(
    minimal_resume_data: dict,
) -> None:
    """A section with nothing in it must not print its heading.

    Both tracked fixtures carry education, so the education branch was only
    ever exercised in the truthy direction -- and there was no guard on it to
    exercise. minimal_resume_data is the one fixture with every optional
    section empty, and it had never been handed to the builder at all: it was
    used for model validation and for the HTTP endpoint, both of which are
    happy with a bare heading.
    """
    resume = Resume.model_validate(minimal_resume_data)
    assert not resume.projects
    assert not resume.education
    assert not resume.credentials

    texts = paragraph_texts(build_resume_document(resume))
    for heading in ("PROJECTS", "EDUCATION", "CREDENTIALS"):
        assert heading not in texts, f"bare heading with no content: {heading}"
    for heading in ("SUMMARY", "SKILLS", "EXPERIENCE"):
        assert heading in texts, f"required section missing: {heading}"


def test_no_word_heading_styles_are_used(resume: Resume) -> None:
    """Explicit formatting only.

    Word's built-in heading styles carry theme-dependent formatting that
    renders inconsistently across Word versions and platforms, so the builder
    styles headings itself. Reintroducing them is called out in CLAUDE.md as a
    thing not to do.
    """
    doc = build_resume_document(resume)
    used = {p.style.name for p in doc.paragraphs if p.style is not None}
    assert not any(name.startswith("Heading") for name in used), used
    assert "Title" not in used


def test_section_headings_keep_with_next(resume: Resume) -> None:
    """A heading must never strand at the foot of a page."""
    doc = build_resume_document(resume)
    headings = set(SECTION_HEADINGS)
    found = [p for p in doc.paragraphs if p.text in headings]
    assert found, "no section headings found"
    for p in found:
        assert p.paragraph_format.keep_with_next is True, (
            f"heading not bound to its content: {p.text}"
        )


def test_bullets_are_free_to_break(resume: Resume) -> None:
    """The deliberate non-invariant.

    Widow/orphan protection covers the header chain only. A bullet list that
    refuses to split wastes more page than the orphan it prevents, so bullets
    must not carry keep_with_next.
    """
    doc = build_resume_document(resume)
    bullets = [
        p
        for p in doc.paragraphs
        if p.style is not None and p.style.name == "List Bullet"
    ]
    assert bullets, "no bullets rendered"
    for p in bullets:
        assert p.paragraph_format.keep_with_next is not True, (
            f"bullet pinned to next paragraph: {p.text[:60]}"
        )


def test_page_margins_are_configured(resume: Resume) -> None:
    doc = build_resume_document(resume)
    section = doc.sections[0]
    assert section.top_margin == section.bottom_margin
    assert section.left_margin == section.right_margin
    assert section.left_margin > 0


def test_render_is_deterministic(resume: Resume) -> None:
    """Same input, same document.

    Not a byte comparison -- python-docx embeds timestamps -- but the visible
    content must not vary between runs.
    """
    first = document_text(build_resume_document(resume))
    second = document_text(build_resume_document(resume))
    assert first == second


# ---------------------------------------------------------------------------
# The section manifest
#
# These are the tests that make the manifest real rather than decorative. The
# resume's shape used to be fixed in three files at once -- the JSON schema's
# required keys, this builder's fixed call sequence, and the heading strings
# written into each renderer's body -- so renaming EDUCATION or moving
# CREDENTIALS above EXPERIENCE was a code change in Go, Python, and a schema.
# Each test below is one of those three things becoming a row.
# ---------------------------------------------------------------------------


def headings_in_order(doc: DocumentObject, candidates: set[str]) -> list[str]:
    """The section headings actually printed, in document order."""
    return [text for text in paragraph_texts(doc) if text in candidates]


def test_manifest_heading_is_what_prints(minimal_resume_data: dict) -> None:
    """The heading is the user's text, not a constant in this file.

    This is the whole point of `heading` being separate from `key`: the same
    education block prints under "EDUCATION", "EDUCATION & TRAINING", or
    "FORMATION" depending on whose resume it is.
    """
    minimal_resume_data["education"] = [{"institution": "State University"}]
    minimal_resume_data["sections"] = [
        {"key": "education", "heading": "EDUCATION & TRAINING"}
    ]

    doc = build_resume_document(Resume.model_validate(minimal_resume_data))
    texts = paragraph_texts(doc)

    assert "EDUCATION & TRAINING" in texts
    assert "EDUCATION" not in texts


def test_manifest_order_is_document_order(resume_data: dict) -> None:
    """Reordering the manifest reorders the document.

    Asserted against a reversal rather than a single swap, so a renderer that
    kept its old fixed sequence and merely relabelled cannot pass.
    """
    reversed_sections = list(reversed(resume_data["sections"]))
    resume_data["sections"] = reversed_sections

    doc = build_resume_document(Resume.model_validate(resume_data))

    wanted = [s["heading"] for s in reversed_sections]
    printed = headings_in_order(doc, set(wanted))
    # Sections with no content print nothing at all, so the expected list is
    # the manifest filtered down to what actually appeared -- order is what is
    # under test here, not presence.
    assert printed == [h for h in wanted if h in printed]
    assert printed, "no sections printed at all"
    assert printed != [h for h in SECTION_HEADINGS if h in printed]


def test_section_absent_from_manifest_does_not_print(resume_data: dict) -> None:
    """A hidden section reaches the renderer as an absent one.

    Generation drops `hidden` rows rather than passing a flag, so this is the
    renderer half of that contract: what is not in the manifest is not in the
    document, even though the content block is still right there in the JSON.
    """
    resume_data["sections"] = [
        s for s in resume_data["sections"] if s["key"] != "skills"
    ]

    doc = build_resume_document(Resume.model_validate(resume_data))

    assert "SKILLS" not in paragraph_texts(doc)
    # The content is still in the document object; only the section is gone.
    assert Resume.model_validate(resume_data).skills is not None
    # And the rest of the resume is untouched.
    assert "EXPERIENCE" in paragraph_texts(doc)


def test_unknown_section_key_is_skipped_not_fatal(resume_data: dict) -> None:
    """A key this renderer does not know must not take the document down.

    Otherwise adding a section type becomes a lockstep deploy: a document
    naming PUBLICATIONS would 500 against every renderer that had not shipped
    yet, instead of rendering everything else.
    """
    resume_data["sections"] = [
        {"key": "publications", "heading": "PUBLICATIONS"},
        *resume_data["sections"],
    ]

    doc = build_resume_document(Resume.model_validate(resume_data))

    assert "PUBLICATIONS" not in paragraph_texts(doc)
    assert "EXPERIENCE" in paragraph_texts(doc)


def test_empty_manifest_falls_back_to_the_conventional_order(
    resume_data: dict,
) -> None:
    """An empty manifest means "no manifest", not "print nothing".

    A document produced before manifests existed carries no `sections` key, and
    Pydantic defaults it to []. Reading that as "the user turned every section
    off" would render a page with a name on it and nothing else -- a silent,
    total data loss that still returns 200.
    """
    resume_data["sections"] = []

    doc = build_resume_document(Resume.model_validate(resume_data))
    texts = paragraph_texts(doc)

    assert "EXPERIENCE" in texts
    assert "SUMMARY" in texts


def test_identity_prints_without_being_in_the_manifest(resume_data: dict) -> None:
    """Identity is the document header, not a section.

    There is no manifest row that can turn a resume's name off, and there
    should not be one.
    """
    resume_data["sections"] = []
    identity_name = resume_data["identity"]["name"]

    doc = build_resume_document(Resume.model_validate(resume_data))

    assert identity_name in document_text(doc)
