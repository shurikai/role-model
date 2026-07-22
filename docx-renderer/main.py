from datetime import datetime
import re
from fastapi import FastAPI, Response
from models import Resume
from docx import Document
import io

app = FastAPI()

### Status:
# FORMATTING TODO:
# - join email/phone/links into single contact line, separator "·"
# - reconsider name heading level / spacing before headline

@app.post("/render")
def render(resume: Resume) -> Response:
    # TODO: Refactor document generation into a service layer
    doc = Document()
    doc.add_heading(resume.identity.name, level=0)

    if resume.identity.headline:
        doc.add_paragraph(resume.identity.headline)

    if resume.identity.location:
        doc.add_paragraph(resume.identity.location)

    contact = resume.identity.email
    if resume.identity.phone:
        contact += " · " + resume.identity.phone
    doc.add_paragraph(contact)

    if resume.identity.linkedin_url:
        doc.add_paragraph(resume.identity.linkedin_url)
    if resume.identity.github_url:
        doc.add_paragraph(resume.identity.github_url)
    if resume.identity.site_url:
        doc.add_paragraph(resume.identity.site_url)

    ### Summary section
    doc.add_heading("SUMMARY", level=1)
    doc.add_paragraph(resume.summary)

    ### Skills section
    doc.add_heading("SKILLS", level=1)
    for category, skills_list in resume.skills.items():
        if skills_list:
            skills = ", ".join(skills_list)
            line = f"{category}: {skills}"
            doc.add_paragraph(line, style="List Bullet")

    ### Work Experience section
    doc.add_heading("EXPERIENCE", level=1)
    for employer_block in resume.experience:
        doc.add_heading(employer_block.employer, level=2)

        ### One employer
        for position in employer_block.positions:
            doc.add_heading(position.title, level=3)
            started = format_date(position.started_on)
            ended = format_date(position.ended_on) if position.ended_on else "Present"
            doc.add_paragraph(f"{started} – {ended}")

            for bullet in position.bullets:
                doc.add_paragraph(bullet.text, style="List Bullet")

    ### Projects
    if resume.projects:
        doc.add_heading("PROJECTS", level=1)
        for project in resume.projects:
            p = doc.add_paragraph()
            run = p.add_run(project.name)
            run.bold = True

            if project.tagline:
                p.add_run(f" - {project.tagline}")

            for bullet in project.bullets:
                doc.add_paragraph(bullet.text, "List Bullet")

            if project.writeup_url:
                doc.add_paragraph(project.writeup_url)
            if project.repo_url:
                doc.add_paragraph(project.repo_url)
            if project.live_url:
                doc.add_paragraph(project.live_url)

    ### Education
    ### TODO: Simple naive render of education for now. This needs a refactor to improve.
    doc.add_heading("EDUCATION", level=1)
    for education_block in resume.education:
        # TODO: This is a known bug. If any of these fields is None in the source data, it will be rendered as 'None' in 
        # the output document. Come back here and guard each of these fields when refactoring.
        doc.add_paragraph(f"{education_block.institution} - {education_block.degree} - {education_block.field_of_study} - {education_block.graduated}")

    buffer = io.BytesIO()
    doc.save(buffer)
    content = buffer.getvalue()

    media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    headers = { 'Content-Disposition': 'attachment; filename="resume.docx"' }
    response = Response(content=content, status_code=200, headers=headers, media_type=media_type)
    return response

def format_date(value: str) -> str:
    return datetime.strptime(value, "%Y-%m").strftime("%b %Y")
